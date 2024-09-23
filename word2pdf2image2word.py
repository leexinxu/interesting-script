from docx import Document
from pdf2image import convert_from_path
from docx.shared import Inches
import os
import tempfile
import win32com.client as win32

# 将Word文档转换为PDF（使用Microsoft Word应用）
def convert_docx_to_pdf(input_path, output_path):
    # 创建一个 Word 应用实例
    word = win32.Dispatch("Word.Application")
    word.Visible = False  # 不显示 Word 窗口
    
    # 打开 Word 文档
    doc = word.Documents.Open(input_path)
    
    # 将文档保存为 PDF，FileFormat=17 表示 PDF
    doc.SaveAs(output_path, FileFormat=17)
    
    # 关闭文档
    doc.Close()
    
    # 退出 Word 应用
    word.Quit()

# 将PDF的每一页转换为图片
def convert_pdf_to_images(pdf_path):
    # 指定 Poppler 安装的路径
    poppler_path = os.path.join(os.getcwd(), 'poppler-24.07.0', 'Library', 'bin')
    return convert_from_path(pdf_path, poppler_path=poppler_path)

# 将图片插入到新的Word文档，每页最多放 image2word_num 张图片
def images_to_word(images, image2word_num, output_docx):
    doc = Document()
    
    # 根据 image2word_num 动态确定表格的行列数
    rows = cols = int(image2word_num**0.5)  # 取行和列为 image2word_num 的平方根
    if rows * cols < image2word_num:  # 如果 image2word_num 不是完全平方数，可能需要增加一列
        cols += 1

    # 创建一个临时目录用于保存图片文件
    with tempfile.TemporaryDirectory() as tmpdirname:
        total_images = len(images)
        
        for i in range(0, total_images, image2word_num):
            # 每页创建固定 rows x cols 表格
            table = doc.add_table(rows=rows, cols=cols)
            
            for row in range(rows):
                for col in range(cols):
                    index = i + row * cols + col
                    if index < total_images:
                        cell = table.cell(row, col)
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()
                        
                        # 将 Pillow 图片保存到临时文件
                        image_path = os.path.join(tmpdirname, f'image_{index}.png')
                        images[index].save(image_path)
                        
                        # 调整图片大小适应单元格
                        run.add_picture(image_path, width=Inches(3))
                    else:
                        # 没有图片时插入空白单元格，保持布局
                        table.cell(row, col).text = ""
            
            # 添加分页，除非是最后一页
            if i + image2word_num < total_images:
                doc.add_page_break()
    
    # 保存新的Word文档
    doc.save(output_docx)

def get_input_or_default(prompt, default_value):
    """获取用户输入，如果用户未输入内容，则返回默认值"""
    user_input = input(f"{prompt} (按回车键使用默认值: {default_value}): ")
    if user_input.strip() == "":
        return default_value
    return os.path.abspath(user_input)

# 主函数
def main():
    try:
        # 提示用户输入路径
        file_path = os.path.abspath(input("请输入DOCX或PDF文件路径（DOCX转PDF依赖Microsoft Word应用，生成DOCX在程序目录下output.docx）："))
        docx_path = os.path.abspath("input.docx")
        pdf_path = os.path.abspath("input.pdf")
        output_docx = os.path.abspath("output.docx")
        if(file_path.endswith(".docx")):
            docx_path = file_path
        elif(file_path.endswith(".pdf")):
            pdf_path = file_path
            docx_path = None
        
        #image2word_num = input("几页图片合并成一页，直接回车默认4（生成DOCX在程序目录下output.docx）：")
        image2word_num = 4
        # 如果用户没有输入内容，则使用默认值 4
        if not image2word_num:
            image2word_num = 4
        else:
            image2word_num = int(image2word_num)
        
        print(f"DOCX文件路径: {docx_path}")
        print(f"PDF文件路径: {pdf_path}")
        print(f"几页图片合并成一页: {image2word_num}")
        print(f"生成的Word文件路径: {output_docx}")

        # 1. 将Word文件转换为PDF
        if docx_path:
            print("将Word文件转换为PDF")
            convert_docx_to_pdf(docx_path, pdf_path)
        
        # 2. 将PDF的每一页转换为图片
        print("将PDF的每一页转换为图片")
        images = convert_pdf_to_images(pdf_path)
        
        # 3. 将图片插入到新的Word文档中，每页最多放4张图片
        print("将图片插入到新的Word文档中")
        images_to_word(images, image2word_num, output_docx)

    except Exception as e:
        # 捕获所有异常，打印异常信息
        print(f"程序发生了异常：{e}")
        
    finally:
        # 无论是否发生异常，都会提示按任意键退出
        input("按任意键退出...")

if __name__ == "__main__":
    main()
